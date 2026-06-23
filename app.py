import streamlit as st
import pandas as pd
import os
import eng_to_ipa as ipa
from gtts import gTTS
import re
import io
from pathlib import Path
import nltk
import pydub
import zipfile

# ────────────────────────────────────────────
# 유틸리티 (수정 및 분석 영역)
# ────────────────────────────────────────────
def sanitize_filename(name):
    """파일명으로 사용할 수 없는 문자 제거"""
    name = str(name)
    return re.sub(r'[\\/*?:"<>|]', "", name).strip()


def get_update_filename(original_name: str) -> str:
    """원본 파일명에 _update를 붙인 이름 반환 (확장자 앞에 삽입)"""
    p = Path(original_name)
    return p.stem + "_update" + p.suffix


def create_zip_of_directory(directory_path, sub_directory=None):
    """지정된 디렉토리(directory_path) 혹은 그 하위의 특정 서브디렉토리(sub_directory) 내의 모든 mp3 파일을 ZIP 바이너리로 압축"""
    if not os.path.exists(directory_path):
        return None
    
    search_path = os.path.join(directory_path, sub_directory) if sub_directory else directory_path
    if not os.path.exists(search_path):
        return None

    zip_buffer = io.BytesIO()
    has_files = False
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for root, dirs, files in os.walk(search_path):
            for file in files:
                if file.endswith('.mp3'):
                    file_path = os.path.join(root, file)
                    relative_path = os.path.relpath(file_path, directory_path)
                    zip_file.write(file_path, relative_path)
                    has_files = True
    return zip_buffer.getvalue() if has_files else None


# ────────────────────────────────────────────
# 처리 함수
# ────────────────────────────────────────────
def check_typos(df):
    """단어 오탈자 점검: F열 단어가 순수 알파벳+공백인지 확인하여 의심 단어 표시"""
    st.info("🔍 단어 오탈자를 점검합니다...")
    suspect = []
    for index, row in df.iterrows():
        val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""
        if val_f and not re.match(r"^[A-Za-z\s\-'\.]+$", val_f):
            suspect.append({"행 번호": index + 2, "F열 단어": val_f})

    if suspect:
        st.warning(f"⚠️ 오탈자 의심 단어 {len(suspect)}개 발견:")
        st.dataframe(pd.DataFrame(suspect))
    else:
        st.success("✅ 오탈자 의심 단어 없음 (모든 단어가 올바른 형식)")
    return df


def update_meaning(df):
    """의미 업데이트: G열에 한국어 의미를 채움 (deep-translator 사용)"""
    st.info("📖 한국어 의미를 번역하고 업데이트합니다...")
    
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='en', target='ko')
    except ImportError:
        st.error("deep-translator 라이브러리가 필요합니다. 'pip install deep-translator'를 실행해주세요.")
        return df

    progress_bar = st.progress(0)
    status_text = st.empty()
    total_rows = len(df)

    # G열(인덱스 6) 확보
    while len(df.columns) < 7:
        df[f"Column_{len(df.columns)}"] = ""

    count = 0
    for index, row in df.iterrows():
        val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""
        if val_f:
            existing_g = str(df.iat[index, 6]).strip() if not pd.isna(df.iat[index, 6]) else ""
            # G열이 비어있거나 이전에 생성된 플레이스홀더([단어])인 경우 업데이트
            if not existing_g or existing_g == "nan" or (existing_g.startswith("[") and existing_g.endswith("]")):
                try:
                    translated = translator.translate(val_f)
                    df.iat[index, 6] = translated
                    count += 1
                except Exception as e:
                    st.warning(f"번역 오류 ({val_f}): {e}")

        progress = (index + 1) / total_rows
        progress_bar.progress(progress)
        status_text.text(f"의미 번역 중: {index + 1}/{total_rows} ({val_f})")

    st.success(f"의미 업데이트 완료: {count}개 행 처리 (한국어 번역 결과 저장)")
    return df


def update_ipa(df, base_output_dir):
    """발음기호 업데이트: H열에 IPA 발음기호 삽입"""
    st.info("🔤 발음 기호를 업데이트합니다...")
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_rows = len(df)

    # H열(인덱스 7) 확보
    while len(df.columns) < 8:
        df[f"Column_{len(df.columns)}"] = ""

    count = 0
    for index, row in df.iterrows():
        val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""
        if val_f:
            pronunciation = ipa.convert(val_f)
            df.iat[index, 7] = pronunciation
            count += 1

        progress = (index + 1) / total_rows
        progress_bar.progress(progress)
        status_text.text(f"IPA 처리 중: {index + 1}/{total_rows} ({val_f})")

    st.success(f"발음기호 업데이트 완료: {count}개 행")
    return df


def update_pos(df):
    """품사 업데이트: I열에 품사 정보 삽입 (nltk 사용)"""
    st.info("🏷️ 품사를 업데이트합니다...")
    
    # nltk 데이터 준비
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        from nltk import pos_tag, word_tokenize
    except Exception as e:
        st.error(f"nltk 라이브러리 호출 중 오류 발생: {e}")
        return df

    # 품사 태그 맵핑 (간소화)
    pos_map = {
        'NN': 'n.', 'NNS': 'n.', 'NNP': 'n.', 'NNPS': 'n.',
        'VB': 'v.', 'VBD': 'v.', 'VBG': 'v.', 'VBN': 'v.', 'VBP': 'v.', 'VBZ': 'v.',
        'JJ': 'adj.', 'JJR': 'adj.', 'JJS': 'adj.',
        'RB': 'adv.', 'RBR': 'adv.', 'RBS': 'adv.',
        'IN': 'prep.',
        'PRP': 'pron.', 'PRP$': 'pron.',
        'CC': 'conj.',
        'CD': 'num.',
        'UH': 'int.'
    }

    progress_bar = st.progress(0)
    status_text = st.empty()
    total_rows = len(df)

    # I열(인덱스 8) 확보
    while len(df.columns) < 9:
        df[f"Column_{len(df.columns)}"] = ""

    count = 0
    for index, row in df.iterrows():
        val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""
        if val_f:
            try:
                tokens = word_tokenize(val_f)
                tags = pos_tag(tokens)
                if tags:
                    raw_tag = tags[0][1]
                    pretty_tag = pos_map.get(raw_tag, raw_tag.lower())
                    df.iat[index, 8] = pretty_tag
                    count += 1
            except:
                pass

        progress = (index + 1) / total_rows
        progress_bar.progress(progress)
        status_text.text(f"품사 처리 중: {index + 1}/{total_rows} ({val_f})")

    st.success(f"품사 업데이트 완료: {count}개 행")
    return df


def process_reading_excel(sheets_dict):
    """리딩 파일의 모든 워크시트('본문', '단어') 처리"""
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='en', target='ko')
    except ImportError:
        st.error("deep-translator 라이브러리가 필요합니다. 'pip install deep-translator'를 실행해주세요.")
        return sheets_dict

    # 1. '본문' 워크시트 처리 (D열 -> E열)
    if '본문' in sheets_dict:
        st.info("📖 '본문' 시트의 해석을 업데이트합니다...")
        df = sheets_dict['본문']
        while len(df.columns) < 5:
            df[f"Column_{len(df.columns)}"] = ""
        
        # E열 데이터타입을 object로 변환하여 float64 에러 방지
        df[df.columns[4]] = df[df.columns[4]].astype(object)
        
        progress_bar = st.progress(0)
        for index, row in df.iterrows():
            val_d = str(df.iat[index, 3]).strip() if not pd.isna(df.iat[index, 3]) else ""
            if val_d:
                # E열이 비어있는 경우에만(NaN이거나 빈 문자열) 번역 수행
                existing_e = df.iat[index, 4]
                if pd.isna(existing_e) or str(existing_e).strip() == "":
                    try:
                        df.iat[index, 4] = translator.translate(val_d)
                    except Exception as e:
                        st.warning(f"본문 번역 오류 (행 {index+2}): {e}")
            progress_bar.progress((index + 1) / len(df))
        st.success("'본문' 시트 처리 완료")

    # 2. '단어' 워크시트 처리 (C열 -> D열)
    if '단어' in sheets_dict:
        st.info("📖 '단어' 시트의 뜻을 업데이트합니다...")
        df = sheets_dict['단어']
        while len(df.columns) < 4:
            df[f"Column_{len(df.columns)}"] = ""
            
        # D열 데이터타입을 object로 변환하여 float64 에러 방지
        df[df.columns[3]] = df[df.columns[3]].astype(object)
        
        progress_bar = st.progress(0)
        for index, row in df.iterrows():
            val_c = str(df.iat[index, 2]).strip() if not pd.isna(df.iat[index, 2]) else ""
            if val_c:
                # D열이 비어있는 경우에만 번역 수행
                existing_d = df.iat[index, 3]
                if pd.isna(existing_d) or str(existing_d).strip() == "":
                    try:
                        df.iat[index, 3] = translator.translate(val_c)
                    except Exception as e:
                        st.warning(f"단어 번역 오류 (행 {index+2}): {e}")
            progress_bar.progress((index + 1) / len(df))
        st.success("'단어' 시트 처리 완료")

    return sheets_dict


def translate_polite(translator, text):
    """정확한 번역을 위해 다른 조작 없이 원본 텍스트를 그대로 번역"""
    text_str = str(text).strip()
    if not text_str:
        return ""
    
    if text_str.endswith(':'):
        return text_str
        
    try:
        return translator.translate(text_str)
    except:
        return ""


def process_listening_excel(sheets_dict):
    """리스닝 파일의 모든 워크시트('본문', '단어') 처리"""
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='en', target='ko')
    except ImportError:
        st.error("deep-translator 라이브러리가 필요합니다. 'pip install deep-translator'를 실행해주세요.")
        return sheets_dict

    # 1. '본문' 워크시트 처리 (D열 -> E열)
    if '본문' in sheets_dict:
        st.info("🎧 '본문' 시트의 해석을 업데이트합니다...")
        df = sheets_dict['본문']
        while len(df.columns) < 5:
            df[f"Column_{len(df.columns)}"] = ""
        
        # E열 데이터타입을 object로 변환하여 float64 에러 방지
        df[df.columns[4]] = df[df.columns[4]].astype(object)
        
        progress_bar = st.progress(0)
        for index, row in df.iterrows():
            val_d = str(df.iat[index, 3]).strip() if not pd.isna(df.iat[index, 3]) else ""
            if val_d:
                existing_e = df.iat[index, 4]
                if pd.isna(existing_e) or str(existing_e).strip() == "":
                    # "Anne: Hello" 혹은 "Sally:" 와 같이 화자가 처음에 오고 콜론이 붙은 경우 처리
                    match = re.match(r"^([A-Za-z0-9\s\-]+:)\s*(.*)", val_d)
                    if match:
                        speaker = match.group(1).strip()
                        body = match.group(2).strip()
                        if body:
                            try:
                                translated_body = translate_polite(translator, body)
                                df.iat[index, 4] = f"{speaker} {translated_body}"
                            except Exception as e:
                                st.warning(f"본문 번역 오류 (행 {index+2}): {e}")
                        else:
                            # 콜론으로만 끝나는 화자명인 경우 그대로 복사
                            df.iat[index, 4] = val_d
                    else:
                        try:
                            # 일반 문장 번역
                            df.iat[index, 4] = translate_polite(translator, val_d)
                        except Exception as e:
                            st.warning(f"본문 번역 오류 (행 {index+2}): {e}")
            else:
                # D 열이 비었을 때는 E 열도 공백으로 둔다.
                df.iat[index, 4] = ""
            progress_bar.progress((index + 1) / len(df))
        st.success("'본문' 시트 처리 완료")

    # 2. '단어' 워크시트 처리 (C열 -> D열)
    if '단어' in sheets_dict:
        st.info("📖 '단어' 시트의 뜻을 업데이트합니다...")
        df = sheets_dict['단어']
        while len(df.columns) < 4:
            df[f"Column_{len(df.columns)}"] = ""
            
        # D열 데이터타입을 object로 변환하여 float64 에러 방지
        df[df.columns[3]] = df[df.columns[3]].astype(object)
        
        progress_bar = st.progress(0)
        for index, row in df.iterrows():
            val_c = str(df.iat[index, 2]).strip() if not pd.isna(df.iat[index, 2]) else ""
            if val_c:
                # D열이 비어있는 경우에만 번역 수행
                existing_d = df.iat[index, 3]
                if pd.isna(existing_d) or str(existing_d).strip() == "":
                    try:
                        df.iat[index, 3] = translator.translate(val_c)
                    except Exception as e:
                        st.warning(f"단어 번역 오류 (행 {index+2}): {e}")
            progress_bar.progress((index + 1) / len(df))
        st.success("'단어' 시트 처리 완료")

    return sheets_dict


def parse_listening_paragraphs(paragraphs, filename_base, current_unit, speakers=None, mode="대화", start_sentence_no=1, create_sound_path=True):
    """리스닝 스크립트(문단 리스트)를 문장별로 분할하여 데이터를 정제"""
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='en', target='ko')
    except ImportError:
        st.error("deep-translator 라이브러리가 필요합니다. 'pip install deep-translator'를 실행해주세요.")
        return []

    # nltk sentence tokenizer 준비
    try:
        nltk.data.find('tokenizers/punkt_tab')
    except LookupError:
        nltk.download('punkt_tab', quiet=True)
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)

    data = []
    sentence_no = start_sentence_no

    # speakers 소문자 리스트화
    speakers_lower = [s.strip().lower() for s in speakers] if speakers else []

    # 대화 형식인 경우에만 기본적으로 대화 간 공백을 생략하는 형식으로 판단
    is_dialogue = (mode == "대화")

    for i, p_text in enumerate(paragraphs):
        p_text = p_text.strip()
        if not p_text:
            continue

        # Unit 정보 업데이트 (예: 'Unit 01. The Sun')
        if p_text.lower().startswith("unit"):
            current_unit = p_text
            sentence_no = 1  # 단원별로 문장 번호 초기화
            continue

        # Topic 문단 여부 판별 (Topic 다음에는 항상 한 줄의 공백 추가)
        is_topic = p_text.lower().startswith("topic")

        # 대화 형식 파싱 (예: "Sally: Hello. How are you?")
        # 화자가 단독으로 오는 경우 ("Sally:")
        if p_text.endswith(":") and re.match(r"^[A-Za-z0-9\s\-]+:$", p_text):
            sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
            data.append({
                "A": filename_base,
                "B": current_unit,
                "C": sentence_no,
                "D": p_text,
                "E": p_text,  # 해석하지 않고 그대로 둠
                "F": sound_path
            })
            sentence_no += 1
            
            # 문단이 끝난 후 공백 행 추가 검사
            if i < len(paragraphs) - 1:
                if is_topic or not is_dialogue:
                    data.append({
                        "A": filename_base,
                        "B": current_unit,
                        "C": sentence_no,
                        "D": "",
                        "E": "",
                        "F": ""
                    })
                    sentence_no += 1
            continue

        # 1. 수동 입력 시 화자 목록이 제공되었고, 첫 단어가 화자 목록에 있는 경우
        is_speaker_found = False
        words = p_text.split()
        if speakers_lower and words:
            first_word_raw = words[0]
            first_word_clean = re.sub(r"[^A-Za-z0-9]", "", first_word_raw).lower()
            if first_word_clean in speakers_lower:
                is_speaker_found = True
                speaker_name = first_word_raw.rstrip(":")
                body_text = p_text[len(first_word_raw):].strip().lstrip(":").strip()
                
                if mode == "대화":
                    # '대화' 형식: 한 화자의 문장이 둘 이상이더라도 하나의 sentence로 한 줄에 표시
                    sent_text = f"{speaker_name}: {body_text}"
                    meaning = translate_polite(translator, body_text)
                    meaning_text = f"{speaker_name}: {meaning}"
                    
                    sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
                    data.append({
                        "A": filename_base,
                        "B": current_unit,
                        "C": sentence_no,
                        "D": sent_text,
                        "E": meaning_text,
                        "F": sound_path
                    })
                    sentence_no += 1
                else:
                    # '발표' 형식: 문장별로 쪼갠다.
                    sentences = nltk.sent_tokenize(body_text)
                    for idx, sent in enumerate(sentences):
                        # 첫 번째 문장에만 화자 이름 결합, 이후 문장은 본문만
                        if idx == 0:
                            sent_text = f"{speaker_name}: {sent}"
                        else:
                            sent_text = sent
                            
                        meaning = translate_polite(translator, sent)
                        if idx == 0:
                            meaning_text = f"{speaker_name}: {meaning}"
                        else:
                            meaning_text = meaning
                        
                        sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
                        data.append({
                            "A": filename_base,
                            "B": current_unit,
                            "C": sentence_no,
                            "D": sent_text,
                            "E": meaning_text,
                            "F": sound_path
                        })
                        sentence_no += 1

        if not is_speaker_found:
            # "Sally: Hello..." 같은 형식 판별
            match = re.match(r"^([A-Za-z0-9\s\-]+:)(.*)", p_text)
            if match:
                speaker = match.group(1).strip()
                body = match.group(2).strip()
                
                if mode == "대화":
                    sent_text = f"{speaker} {body}"
                    meaning = translate_polite(translator, body)
                    meaning_text = f"{speaker} {meaning}"
                    
                    sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
                    data.append({
                        "A": filename_base,
                        "B": current_unit,
                        "C": sentence_no,
                        "D": sent_text,
                        "E": meaning_text,
                        "F": sound_path
                    })
                    sentence_no += 1
                else:
                    sentences = nltk.sent_tokenize(body)
                    for idx, sent in enumerate(sentences):
                        if idx == 0:
                            sent_text = f"{speaker} {sent}"
                        else:
                            sent_text = sent
                            
                        meaning = translate_polite(translator, sent)
                        if idx == 0:
                            meaning_text = f"{speaker} {meaning}"
                        else:
                            meaning_text = meaning
                        
                        sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
                        data.append({
                            "A": filename_base,
                            "B": current_unit,
                            "C": sentence_no,
                            "D": sent_text,
                            "E": meaning_text,
                            "F": sound_path
                        })
                        sentence_no += 1
            else:
                # 일반 서술문
                sentences = nltk.sent_tokenize(p_text)
                for sent in sentences:
                    meaning = translate_polite(translator, sent)
                    
                    sound_path = (f"{sanitize_filename(current_unit)}/{sanitize_filename(filename_base)}_{sentence_no}.mp3" if current_unit else f"Unassigned/{sanitize_filename(filename_base)}_{sentence_no}.mp3") if create_sound_path else ""
                    data.append({
                        "A": filename_base,
                        "B": current_unit,
                        "C": sentence_no,
                        "D": sent,
                        "E": meaning,
                        "F": sound_path
                    })
                    sentence_no += 1

        # 문단이 끝난 후 공백 행 추가 검사
        if i < len(paragraphs) - 1:
            # Topic 다음이거나, 대화 형식이 아닌 경우(독백/설명문) 공백 행 추가
            if is_topic or not is_dialogue:
                data.append({
                    "A": filename_base,
                    "B": current_unit,
                    "C": sentence_no,
                    "D": "",
                    "E": "",
                    "F": ""
                })
                sentence_no += 1

    return data


def update_sound_paths(df, base_output_dir):
    """사운드파일 경로 업데이트: L열에 상대 경로 저장"""
    st.info("📂 사운드 파일 경로를 업데이트합니다...")
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_rows = len(df)

    # L열(인덱스 11) 확보
    while len(df.columns) < 12:
        df[f"Column_{len(df.columns)}"] = ""

    count = 0
    for index, row in df.iterrows():
        val_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
        val_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else "Unassigned"
        val_d = str(row.iloc[3]).strip() if not pd.isna(row.iloc[3]) else ""
        val_e = str(row.iloc[4]).strip() if not pd.isna(row.iloc[4]) else ""
        val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""

        if val_f:
            folder_name = sanitize_filename(val_b)
            filename_parts = [val_a, val_d, val_e, val_f]
            filename = "_".join([p for p in filename_parts if p]) + ".mp3"
            filename = sanitize_filename(filename)
            relative_path = f"{folder_name}/{filename}"

            df.iat[index, 11] = relative_path
            count += 1

        progress = (index + 1) / total_rows
        progress_bar.progress(progress)
        status_text.text(f"경로 업데이트 중: {index + 1}/{total_rows}")

    st.success(f"사운드 파일 경로 업데이트 완료: {count}개 행")
    return df


def generate_sounds(df, base_output_dir):
    """사운드 파일 생성: MP3 파일 생성 (이미 존재하면 건너뜀)"""
    st.info("🎵 사운드 파일을 생성합니다...")
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_rows = len(df)
    generated_count = 0
    skipped_count = 0

    if not os.path.exists(base_output_dir):
        os.makedirs(base_output_dir)

    val_f = ""
    for index, row in df.iterrows():
        try:
            val_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
            val_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else "Unassigned"
            val_d = str(row.iloc[3]).strip() if not pd.isna(row.iloc[3]) else ""
            val_e = str(row.iloc[4]).strip() if not pd.isna(row.iloc[4]) else ""
            val_f = str(row.iloc[5]).strip() if not pd.isna(row.iloc[5]) else ""

            if not val_f:
                continue

            folder_name = sanitize_filename(val_b)
            target_dir = os.path.join(base_output_dir, folder_name)

            filename_parts = [val_a, val_d, val_e, val_f]
            filename = "_".join([p for p in filename_parts if p]) + ".mp3"
            filename = sanitize_filename(filename)
            file_path = os.path.join(target_dir, filename)

            if not os.path.exists(file_path):
                if not os.path.exists(target_dir):
                    os.makedirs(target_dir)
                tts = gTTS(text=val_f, lang='en')
                tts.save(file_path)
                generated_count += 1
            else:
                skipped_count += 1

        except Exception as e:
            st.warning(f"Row {index+2} 처리 중 오류: {e}")

        progress = (index + 1) / total_rows
        progress_bar.progress(progress)
        status_text.text(
            f"파일 생성 중: {index + 1}/{total_rows} ({val_f}) "
            f"- 새 파일: {generated_count}, 건너뜀: {skipped_count}"
        )

    st.success(f"사운드 파일 생성 완료! (새로 생성: {generated_count}, 건너뜀: {skipped_count})")
    return df


# ────────────────────────────────────────────
# Streamlit UI
# ────────────────────────────────────────────
st.set_page_config(page_title="영어 학습 자료 도구 세트", layout="wide")
st.title("🎧 영어 학습 자료 처리 도구")

# CORS 우회를 위한 부모 창 리스너 주입 (Base64 안전 인코딩 방식)
import base64

parent_js_code = """
console.log('Streamlit CORS Bridge Script Running (Base64 version)');

// 최신 영역 정보를 부모 전역에 보관 (CORS 격리 우회용)
window.latestStart = null;
window.latestEnd = null;

// 1. 인풋창 락 주기적 감시 및 처리 (부모 창 컨텍스트이므로 보안 에러 없음)
function lockInputs() {
    try {
        const labels = document.querySelectorAll('label');
        for (let label of labels) {
            const text = label.innerText.trim();
            if (text.startsWith('시작(초) #') || text.startsWith('종료(초) #')) {
                const htmlFor = label.getAttribute('for');
                const parentContainer = label.closest('[data-testid="stNumberInput"]');
                const input = htmlFor ? document.getElementById(htmlFor) : (parentContainer ? parentContainer.querySelector('input[type="number"]') : null);
                if (input && !input.readOnly && document.activeElement !== input) {
                    input.readOnly = true;
                    input.style.pointerEvents = 'none';
                    input.style.backgroundColor = '#161a24';
                    input.style.color = '#a3a8b4';
                }
            }
        }
    } catch(e) {}
}
setInterval(lockInputs, 500);

// 공통 인풋 갱신 유틸리티 함수 (동기/비동기 흐름 제어)
function setInputValue(labelText, val, callback) {
    const labels = document.querySelectorAll('label');
    let targetInput = null;
    for (let label of labels) {
        if (label.innerText.trim() === labelText) {
            const htmlFor = label.getAttribute('for');
            const parentContainer = label.closest('[data-testid="stNumberInput"]');
            targetInput = htmlFor ? document.getElementById(htmlFor) : (parentContainer ? parentContainer.querySelector('input[type="number"]') : null);
            break;
        }
    }
    
    if (targetInput) {
        targetInput.readOnly = false;
        targetInput.style.pointerEvents = 'auto';
        targetInput.focus();
        
        const valueSetter = Object.getOwnPropertyDescriptor(targetInput, 'value')?.set;
        const prototype = Object.getPrototypeOf(targetInput);
        const prototypeValueSetter = Object.getOwnPropertyDescriptor(prototype, 'value')?.set;
        const setter = valueSetter || prototypeValueSetter;
        
        if (setter) {
            setter.call(targetInput, val);
        } else {
            targetInput.value = val;
        }
        
        targetInput.dispatchEvent(new Event('input', { bubbles: true }));
        targetInput.dispatchEvent(new Event('change', { bubbles: true }));
        
        setTimeout(() => {
            const keyDown = new KeyboardEvent('keydown', { key: 'Enter', code: 'Enter', keyCode: 13, which: 13, bubbles: true });
            const keyUp = new KeyboardEvent('keyup', { key: 'Enter', code: 'Enter', keyCode: 13, which: 13, bubbles: true });
            targetInput.dispatchEvent(keyDown);
            targetInput.dispatchEvent(keyUp);
            
            setTimeout(() => {
                targetInput.blur();
                targetInput.readOnly = true;
                targetInput.style.pointerEvents = 'none';
                if (callback) callback();
            }, 50);
        }, 50);
    } else {
        if (callback) callback();
    }
}

// 2. 메시지 수신 리스너 등록
window.addEventListener('message', (event) => {
    const data = event.data;
    if (!data) return;
    
    // 2-A. Wavesurfer 영역 변경 사항 수신 및 캐싱
    if (data.type === 'WAVEFORM_REGION_UPDATE') {
        window.latestStart = data.start;
        window.latestEnd = data.end;
    }
    
    // 2-B. SET 버튼 클릭 시 최신 영역 정보를 인풋에 세팅 요청 처리
    if (data.type === 'SET_INPUT_VALUE_FROM_LATEST') {
        const idx = data.idx;
        const startVal = window.latestStart;
        const endVal = window.latestEnd;
        
        if (startVal === null || startVal === undefined || startVal === "" || endVal === null || endVal === undefined || endVal === "") {
            alert('파형 그래프에서 설정된 영역을 찾을 수 없습니다.');
            return;
        }
        
        // 순차 비동기 갱신
        setInputValue('시작(초) #' + (idx + 1), startVal, () => {
            setTimeout(() => {
                setInputValue('종료(초) #' + (idx + 1), endVal, null);
            }, 200);
        });
    }

    // 2-C. PASTE 버튼 클릭 시 지정 값을 인풋에 세팅 요청 처리
    if (data.type === 'SET_INPUT_VALUE_DIRECT') {
        const idx = data.idx;
        const startVal = data.start;
        const endVal = data.end;
        
        setInputValue('시작(초) #' + (idx + 1), startVal, () => {
            setTimeout(() => {
                setInputValue('종료(초) #' + (idx + 1), endVal, null);
            }, 200);
        });
    }
    
    // 2-D. 클립보드 복사 요청 처리
    if (data.type === 'COPY_CLIPBOARD') {
        const text = data.text;
        // 부모 전역에 마지막 복사 텍스트 보관 (PASTE 릴레이용)
        window.latestClipboard = text;
        if (navigator.clipboard && navigator.clipboard.writeText) {
            navigator.clipboard.writeText(text).catch(err => {
                console.error('Clipboard copy fail:', err);
            });
        } else {
            const textArea = document.createElement('textarea');
            textArea.value = text;
            textArea.style.position = 'fixed';
            document.body.appendChild(textArea);
            textArea.focus();
            textArea.select();
            document.execCommand('copy');
            document.body.removeChild(textArea);
        }
    }

    // 2-E. PASTE 버튼 클릭 시 부모에 저장된 클립보드 텍스트로 인풋 갱신
    if (data.type === 'PASTE_FROM_CLIPBOARD') {
        const idx = data.idx;
        const clipText = window.latestClipboard;
        if (!clipText) {
            alert('붙여넣을 값이 없습니다. 먼저 파형에서 "모두 복사" 버튼을 눌러주세요.');
            return;
        }
        const parts = clipText.split(',');
        if (parts.length === 2) {
            const startVal = parseFloat(parts[0].trim());
            const endVal = parseFloat(parts[1].trim());
            if (!isNaN(startVal) && !isNaN(endVal)) {
                setInputValue('시작(초) #' + (idx + 1), startVal.toFixed(2), () => {
                    setTimeout(() => {
                        setInputValue('종료(초) #' + (idx + 1), endVal.toFixed(2), null);
                    }, 200);
                });
            } else {
                alert('시간 형식이 올바르지 않습니다.');
            }
        } else {
            alert('붙여넣을 데이터 형식이 올바르지 않습니다. (예: 1.25, 4.80)\n먼저 파형에서 "모두 복사" 버튼을 눌러주세요.');
        }
    }
});
"""

encoded_js = base64.b64encode(parent_js_code.encode('utf-8')).decode('utf-8')
inject_js = f'<img src="x" onerror="if(!window.hasStreamlitListener){{window.hasStreamlitListener=true; eval(atob(\'{encoded_js}\'));}}" style="display:none;">'
st.markdown(inject_js, unsafe_allow_html=True)

# allow-same-origin 샌드박스를 이용한 부모 DOM 직접 접근 릴레이
# inject_js(img onerror)가 CSP에 차단될 경우에도 인풋 잠금이 동작하도록 보완
relay_html = """
<script>
(function() {
    var p = window.parent;
    if (p.__wfLockActive) return;
    p.__wfLockActive = true;
    p.setInterval(function() {
        try {
            var labels = p.document.querySelectorAll('label');
            for (var i = 0; i < labels.length; i++) {
                var t = labels[i].innerText.trim();
                if (t.startsWith('시작(초) #') || t.startsWith('종료(초) #')) {
                    var hf = labels[i].getAttribute('for');
                    var pc = labels[i].closest('[data-testid="stNumberInput"]');
                    var inp = hf ? p.document.getElementById(hf) :
                        (pc ? pc.querySelector('input[type="number"]') : null);
                    if (inp && !inp.readOnly && p.document.activeElement !== inp) {
                        inp.readOnly = true;
                        inp.style.pointerEvents = 'none';
                        inp.style.backgroundColor = '#161a24';
                        inp.style.color = '#a3a8b4';
                    }
                }
            }
        } catch(e) {}
    }, 500);
})();
</script>
"""
st.components.v1.html(relay_html, height=1)


tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
    "🔤 단어 파일 처리", 
    "📚 리딩-해석추가", 
    "📑 리딩 스크립트 업로드", 
    "✍️ 리딩 단원 스크립트 추가",
    "🎧 리스닝 스크립트 업로드",
    "📊 리스닝 파일 처리",
    "📝 리스닝 단원 스크립트 추가",
    "🔊 리스닝 사운드 생성",
    "✂️ 리스닝 사운드 편집"
])

# ==========================================
# TAB 1: 단어 파일 처리
# ==========================================
with tab1:
    st.markdown("""
    단어 엑셀 파일을 업로드하고, 실행할 기능을 **체크박스**로 선택한 뒤 **▶ 선택한 작업 실행** 버튼을 클릭하세요.  
    처리된 엑셀 파일은 원본 파일명에 `_update`가 붙어 저장됩니다.
    """)

    uploaded_voca = st.file_uploader("📁 단어 파일 업로더 (.xlsx)", type=["xlsx"], key="voca_uploader")

    if uploaded_voca is not None:
        voca_filename = uploaded_voca.name

        # 세션 상태에 단어 데이터프레임 저장
        if 'voca_df' not in st.session_state or 'voca_loaded_file' not in st.session_state \
                or st.session_state.voca_loaded_file != voca_filename:
            st.session_state.voca_df = pd.read_excel(uploaded_voca, header=0)
            st.session_state.voca_loaded_file = voca_filename
            st.info(f"✅ '{voca_filename}' 파일이 로드되었습니다.")

        v_df = st.session_state.voca_df

        st.write("### 📊 데이터 미리보기 (하단 5행)")
        st.dataframe(v_df.tail())

        base_output_dir = "output_sounds"

        # ── 체크박스 선택 영역 ──────────────────────
        st.write("---")
        st.write("### ⚙️ 실행할 작업 선택")

        col_left, col_right = st.columns(2)

        with col_left:
            do_typo    = st.checkbox("🔍 단어 오탈자 점검 (F열)", value=False, key="chk_typo")
            do_meaning  = st.checkbox("📖 의미 업데이트 (G열)", value=False, key="chk_meaning")
            do_ipa      = st.checkbox("🔤 발음기호 업데이트 (H열)", value=False, key="chk_ipa")

        with col_right:
            do_pos      = st.checkbox("🏷️ 품사 업데이트 (I열)", value=False, key="chk_pos")
            do_path     = st.checkbox("📂 사운드파일 경로 업데이트 (L열)", value=False, key="chk_path")
            do_sound    = st.checkbox("🎵 사운드 파일 생성 (MP3)", value=False, key="chk_sound")

        st.write("---")

        # ── 프로세스 버튼 ─────────────────────────
        if st.button("▶ 선택한 작업 실행", type="primary", use_container_width=True, key="btn_voca"):
            if not any([do_typo, do_meaning, do_ipa, do_pos, do_path, do_sound]):
                st.warning("⚠️ 실행할 작업을 하나 이상 선택해주세요.")
            else:
                with st.spinner("작업을 수행하는 중입니다..."):
                    current_v_df = st.session_state.voca_df.copy()

                    if do_typo:
                        st.write("#### 1️⃣ 단어 오탈자 점검")
                        current_v_df = check_typos(current_v_df)

                    if do_meaning:
                        st.write("#### 2️⃣ 의미 업데이트")
                        current_v_df = update_meaning(current_v_df)

                    if do_ipa:
                        st.write("#### 3️⃣ 발음기호 업데이트")
                        current_v_df = update_ipa(current_v_df, base_output_dir)

                    if do_pos:
                        st.write("#### 4️⃣ 품사 업데이트")
                        current_v_df = update_pos(current_v_df)

                    if do_path:
                        st.write("#### 5️⃣ 사운드파일 경로 업데이트")
                        current_v_df = update_sound_paths(current_v_df, base_output_dir)

                    if do_sound:
                        st.write("#### 6️⃣ 사운드 파일 생성")
                        current_v_df = generate_sounds(current_v_df, base_output_dir)

                    st.session_state.voca_df = current_v_df

                st.success("✅ 모든 선택 작업이 완료되었습니다!")

        # ── 다운로드 영역 ─────────────────────────
        st.write("---")
        st.write("### 📥 결과 엑셀 파일 다운로드")

        v_output = io.BytesIO()
        with pd.ExcelWriter(v_output, engine='openpyxl') as writer:
            st.session_state.voca_df.to_excel(writer, index=False)
        v_processed_data = v_output.getvalue()

        v_save_filename = get_update_filename(voca_filename)

        st.download_button(
            label=f"💾 '{v_save_filename}' 다운로드",
            data=v_processed_data,
            file_name=v_save_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="dl_voca"
        )

        st.info(f"🔊 사운드 파일 저장 위치: `{os.path.abspath(base_output_dir)}`")


# ==========================================
# TAB 2: 리딩 파일 처리
# ==========================================
with tab2:
    st.markdown("""
    리딩 엑셀 파일을 업로드하세요.  
    **D열**의 영어 문장을 읽어 **E열**이 비어있는 경우 한글로 해석하여 채워줍니다.
    """)

    # ── 템플릿 다운로드 기능 ──────────────────
    st.write("### 📥 리딩 파일 템플릿 다운로드")
    
    # 템플릿 데이터 생성
    template_main_cols = [
        "READING_BOOK_NM", "READING_UNIT_NM", "READING_SENTENCE_SEQ", 
        "READING_SENTENCE", "READING_SETENCE_MEANING", "READING_SENTENCE_SOUND_FILE"
    ]
    template_voca_cols = [
        "READING_BOOK_NM", "READING_UNIT_NM", 
        "READING_VOCA", "READING_VOCA_MEANING", "READING_VOCA_SOUND_FILE"
    ]
    template_links_cols = [
        "title", "link", "teacher_only"
    ]
    
    df_main_template = pd.DataFrame(columns=template_main_cols)
    df_voca_template = pd.DataFrame(columns=template_voca_cols)
    df_links_template = pd.DataFrame(columns=template_links_cols)
    
    template_buffer = io.BytesIO()
    with pd.ExcelWriter(template_buffer, engine='openpyxl') as writer:
        df_main_template.to_excel(writer, sheet_name='본문', index=False)
        df_voca_template.to_excel(writer, sheet_name='단어', index=False)
        df_links_template.to_excel(writer, sheet_name='links', index=False)
    
    st.download_button(
        label="📄 리딩 파일 템플릿(.xlsx) 다운로드",
        data=template_buffer.getvalue(),
        file_name="reading_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="btn_template_download"
    )

    st.write("---")

    uploaded_reading = st.file_uploader("📁 리딩 파일 업로더 (.xlsx)", type=["xlsx"], key="reading_uploader")

    if uploaded_reading is not None:
        reading_filename = uploaded_reading.name

        # 세션 상태에 리딩 데이터프레임(딕셔너리) 저장
        if 'reading_sheets' not in st.session_state or 'reading_loaded_file' not in st.session_state \
                or st.session_state.reading_loaded_file != reading_filename:
            # 모든 시트 로드
            st.session_state.reading_sheets = pd.read_excel(uploaded_reading, sheet_name=None, header=0)
            st.session_state.reading_loaded_file = reading_filename
            st.info(f"✅ '{reading_filename}' 파일(모든 시트)이 로드되었습니다.")

        r_sheets = st.session_state.reading_sheets

        st.write("### 📊 데이터 미리보기 (시트 선택)")
        selected_sheet = st.selectbox("미리보기할 시트를 선택하세요:", list(r_sheets.keys()), key="sheet_selector")
        # 마지막 5개 컬럼만 미리보기에 표시
        st.dataframe(r_sheets[selected_sheet].tail(5))

        st.write("---")

        if st.button("▶ 모든 시트 해석/뜻 업데이트 실행", type="primary", use_container_width=True, key="btn_reading"):
            with st.spinner("번역 작업을 수행하는 중입니다..."):
                # 모든 시트 처리
                processed_sheets = process_reading_excel(st.session_state.reading_sheets.copy())
                st.session_state.reading_sheets = processed_sheets
            st.success("✅ 모든 시트의 업데이트가 완료되었습니다!")

        # ── 다운로드 영역 ─────────────────────────
        st.write("---")
        st.write("### 📥 결과 엑셀 파일 다운로드")

        r_output = io.BytesIO()
        with pd.ExcelWriter(r_output, engine='openpyxl') as writer:
            for sheet_name, sheet_df in st.session_state.reading_sheets.items():
                sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
        r_processed_data = r_output.getvalue()

        r_save_filename = get_update_filename(reading_filename)

        st.download_button(
            label=f"💾 '{r_save_filename}' 다운로드",
            data=r_processed_data,
            file_name=r_save_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="dl_reading"
        )

# ==========================================
# TAB 3: 리딩 스크립트 업로드 (Word -> Excel)
# ==========================================
with tab3:
    st.markdown("""
    워드(.docx) 파일을 업로드하면 문장별로 분리하여 엑셀 형식으로 시트를 만들어줍니다.  
    - **Unit**으로 시작하는 문단은 단원명으로 인식합니다.
    - 문단이 바뀔 때 공백 행을 추가합니다.
    - 영어 문장을 한글로 자동 번역하여 E열에 추가합니다.
    """)

    uploaded_docx = st.file_uploader("📁 워드 파일 업로더 (.docx)", type=["docx"], key="docx_uploader")

    if uploaded_docx is not None:
        if st.button("▶ 엑셀 파일로 변환 실행", type="primary", use_container_width=True):
            try:
                from docx import Document
                from deep_translator import GoogleTranslator
                
                doc = Document(uploaded_docx)
                filename_base = Path(uploaded_docx.name).stem
                translator = GoogleTranslator(source='en', target='ko')
                
                # nltk sentence tokenizer 준비
                try:
                    nltk.data.find('tokenizers/punkt_tab')
                except LookupError:
                    nltk.download('punkt_tab', quiet=True)
                try:
                    nltk.data.find('tokenizers/punkt')
                except LookupError:
                    nltk.download('punkt', quiet=True)
                
                data = []
                current_unit = ""
                sentence_no = 1
                
                # 유효한 문단들만 필터링
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for i, p_text in enumerate(paragraphs):
                    # Unit 정보 업데이트 (예: 'Unit 01. The Sun')
                    if p_text.lower().startswith("unit"):
                        current_unit = p_text
                        sentence_no = 1  # 단원별로 문장 번호 초기화
                        # Unit 행은 데이터에 포함시키지 않거나, 필요에 따라 포함 가능
                        # 여기선 단원명으로만 사용하고 다음 문단부터 처리
                        continue
                    
                    # 문장 분리
                    sentences = nltk.sent_tokenize(p_text)
                    for sent in sentences:
                        # 한글 번역
                        try:
                            meaning = translator.translate(sent)
                        except:
                            meaning = ""
                        
                        # 사운드 경로 생성 규칙: Unit/파일명_번호.mp3
                        unit_folder = sanitize_filename(current_unit) if current_unit else "Unassigned"
                        sound_filename = f"{sanitize_filename(filename_base)}_{sentence_no}.mp3"
                        sound_path = f"{unit_folder}/{sound_filename}"
                        
                        data.append({
                            "A": filename_base,
                            "B": current_unit,
                            "C": sentence_no,
                            "D": sent,
                            "E": meaning,
                            "F": sound_path
                        })
                        sentence_no += 1
                    
                    # 문단이 바뀔 때 공백 행 추가 (사용자 요청: D열이 공백인 경우도 문장번호 기재)
                    if i < len(paragraphs) - 1:
                        # 다음 문단이 Unit으로 시작하면 공백행을 생략할 수도 있지만, 명시적 요청에 따라 일단 추가
                        data.append({
                            "A": filename_base,
                            "B": current_unit,
                            "C": sentence_no,
                            "D": "",
                            "E": "",
                            "F": ""
                        })
                        sentence_no += 1
                    
                    progress_bar.progress((i + 1) / len(paragraphs))
                    status_text.text(f"처리 중: {i + 1}/{len(paragraphs)} 문단 완료")
                
                result_df = pd.DataFrame(data)
                # 컬럼명 변경
                result_df.columns = [
                    "파일명", "Unit명", "문장번호", "영어문장", "한글해석", "사운드경로"
                ]
                
                st.success("✅ 변환 완료!")
                st.dataframe(result_df)
                
                # 다운로드 버튼
                docx_output = io.BytesIO()
                with pd.ExcelWriter(docx_output, engine='openpyxl') as writer:
                    result_df.to_excel(writer, index=False, sheet_name="본문")
                
                st.download_button(
                    label="📥 변환된 엑셀 파일 다운로드",
                    data=docx_output.getvalue(),
                    file_name=f"{filename_base}_converted.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="dl_docx"
                )
                
            except Exception as e:
                st.error(f"변환 중 오류 발생: {e}")
                st.exception(e)

# ==========================================
# TAB 4: 리딩 단원 스크립트 추가 (수동 입력)
# ==========================================
with tab4:
    st.markdown("""
    교재 제목, 단원 제목, 그리고 본문 스크립트를 직접 입력하여 엑셀 파일로 변환합니다.
    """)
    
    col1, col2 = st.columns(2)
    with col1:
        manual_book_nm = st.text_input("📚 교재 제목", placeholder="예: Middle School English 1")
    with col2:
        manual_unit_nm = st.text_input("📑 단원 제목", placeholder="예: Unit 1. Nice to Meet You")
        
    manual_body = st.text_area("📝 본문 입력 (여러 줄)", height=350, placeholder="여기에 본문 내용을 붙여넣으세요...")

    if st.button("▶ 입력 완료 및 엑셀 생성", type="primary", use_container_width=True):
        if not manual_book_nm or not manual_unit_nm or not manual_body:
            st.warning("⚠️ 교재 제목, 단원 제목, 본문을 모두 입력해주세요.")
        else:
            with st.spinner("처리 중입니다..."):
                try:
                    from deep_translator import GoogleTranslator
                    translator = GoogleTranslator(source='en', target='ko')
                    
                    # nltk 준비
                    try:
                        nltk.data.find('tokenizers/punkt_tab')
                    except LookupError:
                        nltk.download('punkt_tab', quiet=True)
                    
                    # 본문을 문단별로 나누고 문장으로 분리
                    paragraphs = [p.strip() for p in manual_body.split('\n') if p.strip()]
                    
                    manual_data = []
                    sentence_no = 1
                    
                    for i, p_text in enumerate(paragraphs):
                        sentences = nltk.sent_tokenize(p_text)
                        for sent in sentences:
                            try:
                                meaning = translator.translate(sent)
                            except:
                                meaning = ""
                            
                            manual_data.append({
                                "파일명": manual_book_nm,
                                "Unit명": manual_unit_nm,
                                "문장번호": sentence_no,
                                "영어문장": sent,
                                "한글해석": meaning,
                                "사운드경로": f"{sanitize_filename(manual_unit_nm)}/{sanitize_filename(manual_book_nm)}_{sentence_no}.mp3"
                            })
                            sentence_no += 1
                        
                        # 문단 간 공백 행 추가
                        if i < len(paragraphs) - 1:
                            manual_data.append({
                                "파일명": manual_book_nm,
                                "Unit명": manual_unit_nm,
                                "문장번호": sentence_no,
                                "영어문장": "",
                                "한글해석": "",
                                "사운드경로": ""
                            })
                            sentence_no += 1
                            
                    manual_df = pd.DataFrame(manual_data)
                    st.success("✅ 변환이 완료되었습니다!")
                    st.dataframe(manual_df)
                    
                    # 다운로드
                    m_output = io.BytesIO()
                    with pd.ExcelWriter(m_output, engine='openpyxl') as writer:
                        manual_df.to_excel(writer, index=False, sheet_name="본문")
                    
                    st.download_button(
                        label="📥 생성된 엑셀 파일 다운로드",
                        data=m_output.getvalue(),
                        file_name=f"{manual_unit_nm}_script.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="dl_manual"
                    )
                except Exception as e:
                    st.error(f"오류 발생: {e}")


# ==========================================
# TAB 5: 리스닝 스크립트 업로드
# ==========================================
with tab5:
    st.markdown("""
    워드(.docx) 파일을 업로드하면 대화 형식 및 문장을 분리하여 리스닝 엑셀 형식으로 시트를 만들어줍니다.  
    - **Unit**으로 시작하는 문단은 단원명으로 인식합니다.
    - 대화 형식(`Sally: Hello. How are you?`)은 화자 이름과 함께 문장별로 쪼개어 처리합니다.
    - 한 명이 여러 문장을 말한 경우 각각의 문장이 분리됩니다.
    - 문단이 바뀔 때는 D 열에 공백 행을 추가합니다. (C 열 문장번호는 유지)
    - 영어 문장을 한글로 자동 번역하여 E 열에 추가합니다.
    """)

    uploaded_listen_docx = st.file_uploader("📁 리스닝 워드 파일 업로더 (.docx)", type=["docx"], key="listen_docx_uploader")

    if uploaded_listen_docx is not None:
        if st.button("▶ 리스닝 엑셀 파일로 변환 실행", type="primary", use_container_width=True, key="btn_listen_docx_convert"):
            try:
                from docx import Document
                
                doc = Document(uploaded_listen_docx)
                filename_base = Path(uploaded_listen_docx.name).stem
                
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                with st.spinner("리스닝 워드 파일을 분석하고 번역하는 중..."):
                    data = parse_listening_paragraphs(paragraphs, filename_base, "")
                
                if data:
                    result_df = pd.DataFrame(data)
                    result_df.columns = ["파일명", "Unit명", "문장번호", "영어문장", "한글해석", "사운드경로"]
                    
                    st.success("✅ 변환 완료!")
                    st.dataframe(result_df)
                    
                    # 다운로드 버튼
                    docx_output = io.BytesIO()
                    with pd.ExcelWriter(docx_output, engine='openpyxl') as writer:
                        result_df.to_excel(writer, index=False, sheet_name="본문")
                    
                    st.download_button(
                        label="📥 변환된 리스닝 엑셀 파일 다운로드",
                        data=docx_output.getvalue(),
                        file_name=f"{filename_base}_listening_converted.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        key="dl_listen_docx"
                    )
                else:
                    st.warning("⚠️ 추출된 텍스트 데이터가 없습니다.")
                    
            except Exception as e:
                st.error(f"변환 중 오류 발생: {e}")
                st.exception(e)


# ==========================================
# TAB 6: 리스닝 파일 처리
# ==========================================
with tab6:
    st.markdown("""
    리스닝 엑셀 파일을 업로드하세요.  
    **D열**의 영어 문장을 읽어 **E열**이 비어있는 경우 한글로 해석하여 채워줍니다.  
    단, `"Sally:"`와 같이 콜론으로 끝나는 화자 표시 텍스트는 번역하지 않고 그대로 유지합니다.
    """)

    # ── 템플릿 다운로드 기능 ──────────────────
    st.write("### 📥 리스닝 파일 템플릿 다운로드")
    
    template_listen_main_cols = [
        "LISTENING_BOOK_NM", "LISTENING_UNIT_NM", "LISTENING_SENTENCE_SEQ", 
        "LISTENING_SENTENCE", "LISTENING_SETENCE_MEANING", "LISTENING_SENTENCE_SOUND_FILE"
    ]
    template_listen_voca_cols = [
        "LISTENING_BOOK_NM", "LISTENING_UNIT_NM", 
        "LISTENING_VOCA", "LISTENING_VOCA_MEANING", "LISTENING_VOCA_SOUND_FILE"
    ]
    template_listen_links_cols = [
        "title", "link", "teacher_only"
    ]
    
    df_listen_main_template = pd.DataFrame(columns=template_listen_main_cols)
    df_listen_voca_template = pd.DataFrame(columns=template_listen_voca_cols)
    df_listen_links_template = pd.DataFrame(columns=template_listen_links_cols)
    
    listen_template_buffer = io.BytesIO()
    with pd.ExcelWriter(listen_template_buffer, engine='openpyxl') as writer:
        df_listen_main_template.to_excel(writer, sheet_name='본문', index=False)
        df_listen_voca_template.to_excel(writer, sheet_name='단어', index=False)
        df_listen_links_template.to_excel(writer, sheet_name='links', index=False)
    
    st.download_button(
        label="📄 리스닝 파일 템플릿(.xlsx) 다운로드",
        data=listen_template_buffer.getvalue(),
        file_name="listening_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="btn_listen_template_download"
    )

    st.write("---")

    uploaded_listening = st.file_uploader("📁 리스닝 파일 업로더 (.xlsx)", type=["xlsx"], key="listening_uploader")

    if uploaded_listening is not None:
        listening_filename = uploaded_listening.name

        # 세션 상태에 리스닝 데이터프레임(딕셔너리) 저장
        if 'listening_sheets' not in st.session_state or 'listening_loaded_file' not in st.session_state \
                or st.session_state.listening_loaded_file != listening_filename:
            st.session_state.listening_sheets = pd.read_excel(uploaded_listening, sheet_name=None, header=0)
            st.session_state.listening_loaded_file = listening_filename
            st.info(f"✅ '{listening_filename}' 파일(모든 시트)이 로드되었습니다.")

        l_sheets = st.session_state.listening_sheets

        st.write("### 📊 데이터 미리보기 (시트 선택)")
        selected_l_sheet = st.selectbox("미리보기할 시트를 선택하세요:", list(l_sheets.keys()), key="l_sheet_selector")
        st.dataframe(l_sheets[selected_l_sheet].tail(5))

        st.write("---")

        if st.button("▶ 모든 시트 해석/뜻 업데이트 실행", type="primary", use_container_width=True, key="btn_listening_process"):
            with st.spinner("번역 작업을 수행하는 중입니다..."):
                processed_l_sheets = process_listening_excel(st.session_state.listening_sheets.copy())
                st.session_state.listening_sheets = processed_l_sheets
            st.success("✅ 모든 시트의 업데이트가 완료되었습니다!")

        # ── 다운로드 영역 ─────────────────────────
        st.write("---")
        st.write("### 📥 결과 엑셀 파일 다운로드")

        l_output = io.BytesIO()
        with pd.ExcelWriter(l_output, engine='openpyxl') as writer:
            for sheet_name, sheet_df in st.session_state.listening_sheets.items():
                sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
        l_processed_data = l_output.getvalue()

        l_save_filename = get_update_filename(listening_filename)

        st.download_button(
            label=f"💾 '{l_save_filename}' 다운로드",
            data=l_processed_data,
            file_name=l_save_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="dl_listening_processed"
        )


# ==========================================
# TAB 7: 리스닝 단원 스크립트 추가 (수동 입력)
# ==========================================
with tab7:
    st.markdown("""
    교재 제목, 단원 제목, 화자 목록을 공통으로 지정하고,  
    여러 개의 본문 박스에 각각 대화 또는 발표 형식의 스크립트를 입력하여 하나의 리스닝용 엑셀 파일로 변환합니다.
    """)
    
    col1_l, col2_l = st.columns(2)
    with col1_l:
        manual_listen_book_nm = st.text_input("📚 교재 제목", placeholder="예: Middle School Listening 1", key="input_l_book_nm")
    with col2_l:
        manual_listen_unit_nm = st.text_input("📑 단원 제목", placeholder="예: Unit 1. Nice to Meet You", key="input_l_unit_nm")
        
    manual_listen_speakers = st.text_input("👤 화자 목록 (쉼표로 구분)", placeholder="예: Sally, John, Teacher", key="input_l_speakers")
    
    # 본문 박스 개수 관리
    if "num_bodies" not in st.session_state:
        st.session_state.num_bodies = 2
        
    col_btn1, col_btn2 = st.columns([1, 4])
    with col_btn1:
        if st.button("➕ 본문 박스 추가", key="add_body_btn"):
            st.session_state.num_bodies += 1
            st.rerun()
            
    bodies_data = []
    
    for i in range(1, st.session_state.num_bodies + 1):
        st.write("---")
        st.subheader(f"📝 본문 박스 {i}")
        
        default_idx = 0 if i == 1 else 1
        mode_i = st.selectbox(f"형식 선택 (박스 {i})", ["대화", "발표"], index=default_idx, key=f"mode_{i}")
        placeholder_text = f"여기에 {i}번째 대화/발표 내용을 붙여넣으세요..."
        body_text_i = st.text_area(f"본문 입력 {i}", height=250, placeholder=placeholder_text, key=f"input_l_body_{i}")
        
        bodies_data.append((body_text_i, mode_i, i))
        
        if body_text_i:
            st.write(f"🔍 **본문 {i} 미리보기 (화자 강조 및 편집)**")
            # 화자 목록 파싱
            speakers = [s.strip() for s in manual_listen_speakers.split(",") if s.strip()] if manual_listen_speakers else []
            preview_html = body_text_i.replace("\n", "<br>")
            if speakers:
                for sp in speakers:
                    pattern = re.compile(rf"\b({re.escape(sp)})\b", re.IGNORECASE)
                    preview_html = pattern.sub(r'<span style="color:red; font-weight:bold;">\1</span>', preview_html)
            
            lines_count = len(body_text_i.split('\n'))
            iframe_height = min(max(200, lines_count * 24 + 110), 450)
            
            editor_html = f"""
            <div style="margin-bottom: 10px;">
                <button id="apply-btn-{i}" style="padding: 6px 12px; background-color: #ff4b4b; color: white; border: none; border-radius: 4px; cursor: pointer; font-weight: bold; font-family: sans-serif;">
                    ✍️ 줄바꿈 수정사항 본문 {i}에 반영하기
                </button>
            </div>
            <div id="editor-{i}" contenteditable="true" style="border:1px solid #555; padding:15px; border-radius:5px; line-height:1.6; font-family: sans-serif; min-height:100px; max-height: 280px; overflow-y: auto; outline:none; color:inherit; background-color:transparent;">
                {preview_html}
            </div>

            <script>
            document.getElementById('apply-btn-{i}').addEventListener('click', () => {{
                const editor = document.getElementById('editor-{i}');
                const editedText = editor.innerText;
                
                try {{
                    const parentDocs = window.parent.document;
                    const textareas = parentDocs.querySelectorAll('textarea');
                    let targetTextarea = null;
                    
                    for (let ta of textareas) {{
                        if (ta.placeholder && ta.placeholder.includes("여기에 {i}번째 대화/발표 내용을 붙여넣으세요")) {{
                            targetTextarea = ta;
                            break;
                        }}
                    }}
                    
                    if (targetTextarea) {{
                        const nativeTextAreaValueSetter = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, "value").set;
                        nativeTextAreaValueSetter.call(targetTextarea, editedText);
                        const event = new Event('input', {{ bubbles: true }});
                        targetTextarea.dispatchEvent(event);
                        alert('✅ 수정사항이 본문 입력 {i}에 성공적으로 반영되었습니다!');
                    }} else {{
                        navigator.clipboard.writeText(editedText);
                        alert('입력창을 직접 찾지 못해 수정된 텍스트를 클립보드에 복사했습니다. 위의 입력창에 붙여넣어(Ctrl+V) 주세요!');
                    }}
                }} catch (e) {{
                    navigator.clipboard.writeText(editedText);
                    alert('수정된 텍스트가 클립보드에 복사되었습니다. 위의 본문 입력 {i}에 전체 붙여넣기(Ctrl+V) 해주세요!');
                }}
            }});
            </script>
            """
            st.components.v1.html(editor_html, height=iframe_height)
            st.write("---")

    # ── 엑셀 생성 및 병합 ───────────────────────
    if st.button("▶ 리스닝 엑셀 생성", type="primary", use_container_width=True, key="btn_listen_manual_create"):
        any_body_filled = any(text.strip() for text, _, _ in bodies_data)
        if not manual_listen_book_nm or not manual_listen_unit_nm or not any_body_filled:
            st.warning("⚠️ 교재 제목, 단원 제목 및 적어도 하나의 본문 내용을 입력해주세요.")
        else:
            with st.spinner("처리 중입니다..."):
                try:
                    speakers = [s.strip() for s in manual_listen_speakers.split(",") if s.strip()] if manual_listen_speakers else []
                    
                    merged_data = []
                    current_sentence_no = 1
                    
                    filled_bodies = [(text, mode, idx) for text, mode, idx in bodies_data if text.strip()]
                    
                    for idx_f, (body_text, mode, idx) in enumerate(filled_bodies):
                        paragraphs = [p.strip() for p in body_text.split('\n') if p.strip()]
                        data_i = parse_listening_paragraphs(
                            paragraphs, manual_listen_book_nm, manual_listen_unit_nm, 
                            speakers=speakers, mode=mode, start_sentence_no=current_sentence_no,
                            create_sound_path=False
                        )
                        if data_i:
                            merged_data.extend(data_i)
                            valid_nums = [item["C"] for item in data_i if item["C"]]
                            if valid_nums:
                                current_sentence_no = max(valid_nums) + 1
                                
                        # 마지막 본문이 아니라면 구분용 공백 행 삽입
                        if idx_f < len(filled_bodies) - 1:
                            merged_data.append({
                                "A": manual_listen_book_nm,
                                "B": manual_listen_unit_nm,
                                "C": current_sentence_no,
                                "D": "",
                                "E": "",
                                "F": ""
                            })
                            current_sentence_no += 1
                    
                    if merged_data:
                        manual_listen_df = pd.DataFrame(merged_data)
                        manual_listen_df.columns = ["파일명", "Unit명", "문장번호", "영어문장", "한글해석", "사운드경로"]
                        
                        st.success("✅ 변환이 완료되었습니다!")
                        st.dataframe(manual_listen_df)
                        
                        # 다운로드
                        ml_output = io.BytesIO()
                        with pd.ExcelWriter(ml_output, engine='openpyxl') as writer:
                            manual_listen_df.to_excel(writer, index=False, sheet_name="본문")
                        
                        st.download_button(
                            label="📥 생성된 리스닝 엑셀 파일 다운로드",
                            data=ml_output.getvalue(),
                            file_name=f"{manual_listen_unit_nm}_listening_script.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                            key="dl_listen_manual"
                        )
                    else:
                        st.warning("⚠️ 처리할 텍스트 데이터가 생성되지 않았습니다.")
                except Exception as e:
                    st.error(f"오류 발생: {e}")


# ==========================================
# TAB 8: 리스닝 사운드 생성 (TTS)
# ==========================================
with tab8:
    st.markdown("""
    업로드한 엑셀 파일의 **D열(영어문장)**을 읽어 TTS(gTTS)를 통해 MP3 사운드 파일을 생성합니다.  
    - **화자 정보**(예: `Sally: Hello` 에서 `Sally:`)는 사운드로 변환하지 않고 문장만 사운드로 만듭니다.
    - 생성된 파일은 **[지정된 경로]/[A열 교재명]/[B열 단원명]/[C열 파일명].mp3** 구조로 저장됩니다.
    """)

    uploaded_l_sound_excel = st.file_uploader("📁 리스닝 엑셀 파일 업로드 (.xlsx)", type=["xlsx"], key="listening_sound_excel_uploader")
    base_output_dir_t8 = st.text_input("📂 사운드 저장 기본 경로", value="output_sounds", key="l_sound_gen_base_dir")

    if uploaded_l_sound_excel is not None:
        sheets_t8 = pd.read_excel(uploaded_l_sound_excel, sheet_name=None, header=0)
        sheet_names_t8 = list(sheets_t8.keys())
        selected_sheet_t8 = st.selectbox(
            "파싱할 시트를 선택하세요:", 
            sheet_names_t8, 
            index=0 if '본문' not in sheet_names_t8 else sheet_names_t8.index('본문'),
            key="sheet_selector_t8"
        )
        
        df_t8 = sheets_t8[selected_sheet_t8]
        st.dataframe(df_t8.head(5))

        if st.button("▶ TTS 사운드 파일 생성 실행", type="primary", use_container_width=True, key="btn_listening_sound_generate"):
            if len(df_t8.columns) < 4:
                st.error("⚠️ 엑셀 파일에 적어도 4개의 열(A:교재명, B:단원명, C:파일명, D:영어문장)이 존재해야 합니다.")
            else:
                with st.spinner("TTS 사운드 파일을 생성하는 중..."):
                    try:
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        total_rows = len(df_t8)
                        generated_count = 0
                        skipped_count = 0

                        for index, row in df_t8.iterrows():
                            # 데이터 추출
                            val_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                            val_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else "Unassigned"
                            val_c = str(row.iloc[2]).strip() if not pd.isna(row.iloc[2]) else ""
                            val_d = str(row.iloc[3]).strip() if not pd.isna(row.iloc[3]) else ""

                            # 유효성 검사
                            if not val_d or not val_c:
                                skipped_count += 1
                                continue

                            # 화자명 제거 처리 ("Anne: Hello" -> "Hello")
                            match = re.match(r"^([A-Za-z0-9\s\-]+:)\s*(.*)", val_d)
                            if match:
                                text_to_speak = match.group(2).strip()
                            else:
                                text_to_speak = val_d

                            # 화자명만 단독으로 있거나 텍스트가 없는 경우 스킵
                            if not text_to_speak:
                                skipped_count += 1
                                continue

                            # 경로 설정
                            target_dir = os.path.join(
                                base_output_dir_t8, 
                                sanitize_filename(val_a) if val_a else "Unassigned", 
                                sanitize_filename(val_b)
                            )
                            filename = sanitize_filename(val_c) + ".mp3"
                            file_path = os.path.join(target_dir, filename)

                            # 저장 폴더 생성 및 TTS 저장
                            if not os.path.exists(target_dir):
                                os.makedirs(target_dir, exist_ok=True)

                            tts = gTTS(text=text_to_speak, lang='en')
                            tts.save(file_path)
                            generated_count += 1

                            progress_bar.progress((index + 1) / total_rows)
                            status_text.text(f"진행 중: {index + 1}/{total_rows} (생성: {generated_count}, 건너뜀: {skipped_count})")

                        st.success(f"🎉 TTS 사운드 파일 생성 완료! (생성됨: {generated_count}개, 건너뜀: {skipped_count}개)")
                        st.info(f"🔊 사운드 파일 저장 위치: `{os.path.abspath(base_output_dir_t8)}`")
                        
                        # ZIP 다운로드 지원
                        zip_data = create_zip_of_directory(base_output_dir_t8)
                        if zip_data:
                            st.download_button(
                                label="📥 생성된 사운드 파일 전체 다운로드 (ZIP)",
                                data=zip_data,
                                file_name="listening_sounds_tts.zip",
                                mime="application/zip",
                                use_container_width=True,
                                key="dl_zip_t8"
                            )
                    except Exception as e:
                        st.error(f"오류 발생: {e}")
# ==========================================
# TAB 9: 리스닝 사운드 편집 (Audio Slicing)
# ==========================================
with tab9:
    def rendering_control_buttons(idx):
        html_code = f"""
        <style>
            html, body {{
                margin: 0;
                padding: 0;
                overflow: hidden;
                background-color: transparent;
                height: 100%;
            }}
            .btn-wrapper {{
                display: flex;
                gap: 6px;
                margin-top: 28px;
                height: 38px;
                width: 100%;
                box-sizing: border-box;
            }}
            .action-btn {{
                flex: 1;
                padding: 0 4px;
                border: none;
                border-radius: 4px;
                cursor: pointer;
                font-size: 11px;
                font-weight: bold;
                color: white;
                box-shadow: 0 2px 4px rgba(0,0,0,0.15);
                transition: opacity 0.15s, box-shadow 0.15s;
                display: flex;
                align-items: center;
                justify-content: center;
                gap: 2px;
                height: 100%;
                box-sizing: border-box;
                outline: none;
            }}
            .btn-set {{
                background: linear-gradient(135deg, #00d2ff 0%, #00a2ff 100%);
            }}
            .btn-paste {{
                background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            }}
            .action-btn:hover {{
                opacity: 0.85;
                box-shadow: 0 4px 8px rgba(0,0,0,0.3);
            }}
            .action-btn:active {{
                opacity: 0.7;
                box-shadow: 0 1px 2px rgba(0,0,0,0.1);
            }}
        </style>
        <div class="btn-wrapper">
            <button id="setBtn_{idx}" class="action-btn btn-set">⚙️ SET</button>
            <button id="pasteBtn_{idx}" class="action-btn btn-paste">📋 PASTE</button>
        </div>

        <script>
        // allow-same-origin 샌드박스: window.parent.document 직접 접근으로 Streamlit 인풋 제어
        function setInputValue(labelText, val, callback) {{
            try {{
                var doc = window.parent.document;
                var labels = doc.querySelectorAll('label');
                var targetInput = null;
                for (var i = 0; i < labels.length; i++) {{
                    if (labels[i].innerText.trim() === labelText) {{
                        var hf = labels[i].getAttribute('for');
                        var pc = labels[i].closest('[data-testid="stNumberInput"]');
                        targetInput = hf ? doc.getElementById(hf) :
                            (pc ? pc.querySelector('input[type="number"]') : null);
                        break;
                    }}
                }}
                if (targetInput) {{
                    targetInput.readOnly = false;
                    targetInput.style.pointerEvents = 'auto';
                    targetInput.focus();
                    var proto = Object.getPrototypeOf(targetInput);
                    var desc = Object.getOwnPropertyDescriptor(proto, 'value');
                    var setter = desc && desc.set;
                    if (setter) setter.call(targetInput, val);
                    else targetInput.value = val;
                    targetInput.dispatchEvent(new Event('input', {{ bubbles: true }}));
                    targetInput.dispatchEvent(new Event('change', {{ bubbles: true }}));
                    setTimeout(function() {{
                        targetInput.dispatchEvent(new KeyboardEvent('keydown', {{ key: 'Enter', keyCode: 13, bubbles: true }}));
                        targetInput.dispatchEvent(new KeyboardEvent('keyup', {{ key: 'Enter', keyCode: 13, bubbles: true }}));
                        setTimeout(function() {{
                            targetInput.blur();
                            targetInput.readOnly = true;
                            targetInput.style.pointerEvents = 'none';
                            if (callback) callback();
                        }}, 50);
                    }}, 50);
                }} else {{
                    if (callback) callback();
                }}
            }} catch (e) {{
                console.error('setInputValue error:', e);
                if (callback) callback();
            }}
        }}

        document.getElementById('setBtn_{idx}').addEventListener('click', function() {{
            var start = localStorage.getItem('wf_region_start');
            var end = localStorage.getItem('wf_region_end');
            if (start === null || end === null) {{
                alert('파형 그래프에서 설정된 영역을 찾을 수 없습니다.');
                return;
            }}
            setInputValue('시작(초) #{idx+1}', start, function() {{
                setTimeout(function() {{
                    setInputValue('종료(초) #{idx+1}', end, null);
                }}, 200);
            }});
        }});

        document.getElementById('pasteBtn_{idx}').addEventListener('click', function() {{
            var clip = localStorage.getItem('wf_clipboard');
            if (!clip) {{
                alert('먼저 파형에서 "모두 복사" 버튼을 눌러주세요.');
                return;
            }}
            var parts = clip.split(',');
            if (parts.length === 2) {{
                var sv = parseFloat(parts[0].trim());
                var ev = parseFloat(parts[1].trim());
                if (!isNaN(sv) && !isNaN(ev)) {{
                    setInputValue('시작(초) #{idx+1}', sv.toFixed(2), function() {{
                        setTimeout(function() {{
                            setInputValue('종료(초) #{idx+1}', ev.toFixed(2), null);
                        }}, 200);
                    }});
                }} else {{
                    alert('시간 형식이 올바르지 않습니다.');
                }}
            }} else {{
                alert('붙여넣을 데이터 형식이 올바르지 않습니다.\n먼저 파형에서 "모두 복사"를 눌러주세요.');
            }}
        }});
        </script>
        """
        st.components.v1.html(html_code, height=72, scrolling=False)

    st.markdown("""
    엑셀 파일과 통 오디오 파일(.mp3)을 업로드하여 각 문장별로 사운드를 끊어냅니다.  
    - 재생하면서 문장의 **시작 시간**과 **종료 시간**을 입력하세요.
    - **[미리듣기]** 버튼을 통해 잘라낼 구간의 소리를 미리 들어볼 수 있습니다.
    - 잘라낸 사운드는 **[지정된 경로]/[A열 교재명]/[B열 단원명]/[C열 파일명].mp3**로 저장됩니다. (동일 파일 존재 시 덮어씀)
    """)

    col_e1, col_e2 = st.columns(2)
    with col_e1:
        uploaded_edit_excel = st.file_uploader("📁 리스닝 엑셀 파일 업로드 (.xlsx)", type=["xlsx"], key="listening_edit_excel_uploader")
    with col_e2:
        uploaded_edit_audio = st.file_uploader("🎵 원본 오디오 파일 업로드 (.mp3)", type=["mp3"], key="listening_edit_audio_uploader")

    base_output_dir_t9 = st.text_input("📂 사운드 편집 저장 기본 경로", value="output_sounds", key="l_sound_edit_base_dir")

    # 엑셀과 오디오가 모두 올라왔을 때 활성화
    if uploaded_edit_excel is not None and uploaded_edit_audio is not None:
        try:
            # 1. 엑셀 로드
            sheets_t9 = pd.read_excel(uploaded_edit_excel, sheet_name=None, header=0)
            sheet_names_t9 = list(sheets_t9.keys())
            selected_sheet_t9 = st.selectbox(
                "편집할 시트를 선택하세요:", 
                sheet_names_t9, 
                index=0 if '본문' not in sheet_names_t9 else sheet_names_t9.index('본문'),
                key="sheet_selector_t9"
            )
            df_t9 = sheets_t9[selected_sheet_t9]

            # 2. 오디오 로드 (pydub 사용)
            with st.spinner("오디오 파일을 불러오는 중..."):
                audio_bytes = uploaded_edit_audio.read()
                audio_segment = pydub.AudioSegment.from_file(io.BytesIO(audio_bytes), format="mp3")
                duration_sec = len(audio_segment) / 1000.0
                
            st.success(f"🎵 원본 오디오가 로드되었습니다. (총 길이: {duration_sec:.2f}초)")
            
            # ── 파형 시각화 및 편집 도구 ────────────────
            st.write("### 📊 사운드 진동 파형 (시작/종료점 영역 설정)")
            st.markdown("파형 양 끝의 **하늘색 막대바**를 드래그하여 조절하거나 파형 위를 **드래그**하여 영역(Region)을 설정하세요. 설정한 구간은 자동으로 **반복 재생**됩니다.")
            
            import base64
            audio_base64 = base64.b64encode(audio_bytes).decode("utf-8")
            
            waveform_html = f"""
            <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
            <style>
                body {{
                    margin: 0;
                    padding: 10px;
                    background-color: #0e1117;
                    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                    color: #fafafa;
                }}
                .wave-container {{
                    background: #00162b;
                    border: 1px solid #002d54;
                    border-radius: 8px;
                    padding: 15px;
                    box-shadow: 0 4px 15px rgba(0,0,0,0.4);
                }}
                #waveform {{
                    background: #000f1f;
                    border-radius: 6px;
                    padding: 5px;
                    border: 1px solid #002447;
                }}
                .control-bar {{
                    display: flex;
                    flex-wrap: wrap;
                    align-items: center;
                    justify-content: space-between;
                    gap: 15px;
                    margin-top: 15px;
                    padding-top: 15px;
                    border-top: 1px solid #002d54;
                }}
                .btn-group {{
                    display: flex;
                    gap: 8px;
                }}
                .btn {{
                    padding: 8px 16px;
                    border: none;
                    border-radius: 6px;
                    font-weight: 600;
                    font-size: 13px;
                    cursor: pointer;
                    display: inline-flex;
                    align-items: center;
                    gap: 6px;
                    transition: all 0.2s ease;
                }}
                .btn-primary {{
                    background: linear-gradient(135deg, #ff4b4b 0%, #ff7575 100%);
                    color: white;
                }}
                .btn-secondary {{
                    background: #002447;
                    color: #fafafa;
                    border: 1px solid #003b73;
                }}
                .btn-secondary:not(:disabled) {{
                    background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
                    color: white !important;
                    border: 1px solid #047857 !important;
                    box-shadow: 0 0 8px rgba(16, 185, 129, 0.4);
                }}
                .btn:hover:not(:disabled) {{
                    opacity: 0.85;
                    box-shadow: 0 4px 12px rgba(0,0,0,0.3);
                }}
                .btn:active:not(:disabled) {{
                    opacity: 0.7;
                }}
                .btn:disabled {{
                    background: #001224;
                    color: #475569;
                    border: 1px solid #001c38;
                    cursor: not-allowed;
                }}
                .info-display {{
                    font-size: 14px;
                    color: #94a3b8;
                    background: #001c38;
                    padding: 8px 14px;
                    border-radius: 6px;
                    border: 1px solid #002d54;
                    min-width: 150px;
                    text-align: center;
                }}
                .time-span {{
                    font-weight: bold;
                    color: #00ffd2;
                }}
                .slider-container {{
                    display: flex;
                    align-items: center;
                    gap: 10px;
                    color: #94a3b8;
                    font-size: 13px;
                }}
                .slider-container input[type="range"] {{
                    accent-color: #00d2ff;
                    cursor: pointer;
                }}
                /* Toast 알림 */
                .toast {{
                    position: fixed;
                    bottom: 20px;
                    right: 20px;
                    background: #10b981;
                    color: white;
                    padding: 10px 20px;
                    border-radius: 6px;
                    font-size: 13px;
                    font-weight: bold;
                    box-shadow: 0 4px 10px rgba(0,0,0,0.3);
                    opacity: 0;
                    transition: opacity 0.3s ease;
                    z-index: 9999;
                }}
                .toast.show {{
                    opacity: 1;
                }}
                
                /* Wavesurfer Region 및 Handle Custom Styling */
                .wavesurfer-region {{
                    border: 1px dashed rgba(0, 210, 255, 0.4) !important;
                    background-color: rgba(0, 150, 255, 0.15) !important;
                }}
                .wavesurfer-handle {{
                    width: 6px !important;
                    background-color: #00d2ff !important;
                    opacity: 1 !important;
                    cursor: col-resize;
                    transition: background-color 0.2s;
                }}
                .wavesurfer-handle:hover {{
                    background-color: #00ffd2 !important;
                }}
            </style>
            
            <div class="wave-container">
                <div id="waveform"></div>
                
                <div class="control-bar">
                    <div class="btn-group">
                        <button id="playBtn" class="btn btn-primary"><i class="fa-solid fa-play"></i> 재생 / 일시정지</button>
                    </div>
                    
                    <div class="slider-container">
                        <i class="fa-solid fa-magnifying-glass-minus"></i>
                        <input type="range" id="zoomSlider" min="10" max="300" value="10" style="width: 120px;">
                        <i class="fa-solid fa-magnifying-glass-plus"></i>
                    </div>

                    <div class="info-display">
                        선택 구간: <span id="region-time" class="time-span">없음 (드래그하세요)</span>
                    </div>

                    <div class="btn-group">
                        <button id="copyStartBtn" class="btn btn-secondary" disabled><i class="fa-regular fa-copy"></i> 시작 복사</button>
                        <button id="copyEndBtn" class="btn btn-secondary" disabled><i class="fa-regular fa-copy"></i> 종료 복사</button>
                        <button id="copyBothBtn" class="btn btn-secondary" disabled><i class="fa-solid fa-copy"></i> 모두 복사</button>
                    </div>
                </div>
            </div>
            
            <div id="toast" class="toast">복사되었습니다!</div>

            <script src="https://unpkg.com/wavesurfer.js@7"></script>
            <script src="https://unpkg.com/wavesurfer.js@7/dist/plugins/regions.min.js"></script>
            <script>
                const audioData = "data:audio/mp3;base64,{audio_base64}";
                
                const ws = WaveSurfer.create({{
                    container: '#waveform',
                    waveColor: '#00e1b5',
                    progressColor: '#00d2ff',
                    url: audioData,
                    height: 120,
                    responsive: true
                }});
                
                const wsRegions = ws.registerPlugin(WaveSurfer.Regions.create());
                
                let activeRegion = null;
                let isInitialized = false;

                function initRegionAndUI() {{
                    if (isInitialized) return;
                    try {{
                        const duration = ws.getDuration();
                        if (!duration || isNaN(duration)) return;
                        
                        isInitialized = true;
                        
                        // 오디오 로딩 완료 시 가장 좌측(0초)부터 가장 우측(끝)까지 꽉 찬 region 생성
                        activeRegion = wsRegions.addRegion({{
                            start: 0,
                            end: duration,
                            color: 'rgba(255, 75, 75, 0.25)',
                            drag: true,
                            resize: true
                        }});
                        updateUI(activeRegion);
                        
                        // 부모 리스너의 초기화 지연을 감안하여 500ms, 1500ms 후에 최신 영역값을 안전하게 재전송
                        setTimeout(() => {{
                            if (activeRegion) updateUI(activeRegion);
                        }}, 500);
                        setTimeout(() => {{
                            if (activeRegion) updateUI(activeRegion);
                        }}, 1500);
                        
                        // 드래그 선택 기능도 활성화 (영역을 다 지우고 새로 그릴 때 대비)
                        wsRegions.enableDragSelection({{
                            color: 'rgba(255, 75, 75, 0.25)',
                        }});
                    }} catch (e) {{
                        console.error("Initialization error in initRegionAndUI:", e);
                    }}
                }}

                // decode와 ready 이벤트 모두에 바인딩하여 확실하게 1회 초기화 진행
                ws.on('decode', () => {{
                    initRegionAndUI();
                }});

                ws.on('ready', () => {{
                    initRegionAndUI();
                }});
                
                // 재생/일시정지 제어
                document.getElementById('playBtn').addEventListener('click', () => {{
                    if (activeRegion && !ws.isPlaying()) {{
                        const curr = ws.getCurrentTime();
                        if (curr < activeRegion.start || curr > activeRegion.end) {{
                            ws.setTime(activeRegion.start);
                        }}
                    }}
                    ws.playPause();
                }});
                
                // 줌 제어
                document.getElementById('zoomSlider').addEventListener('input', (e) => {{
                    ws.zoom(Number(e.target.value));
                }});
                
                // 반복 재생 제어 (timeupdate 활용)
                ws.on('timeupdate', () => {{
                    if (activeRegion && ws.isPlaying()) {{
                        const curr = ws.getCurrentTime();
                        if (curr >= activeRegion.end || curr < activeRegion.start) {{
                            ws.setTime(activeRegion.start);
                        }}
                    }}
                }});
                
                wsRegions.on('region-created', (region) => {{
                    // 단 하나의 영역만 유지
                    wsRegions.getRegions().forEach(r => {{
                        if (r !== region) r.destroy();
                    }});
                    activeRegion = region;
                    updateUI(region);
                }});
                
                wsRegions.on('region-updated', (region) => {{
                    activeRegion = region;
                    updateUI(region);
                    
                    // 조절 중에 재생 바가 영역 밖으로 나가면 영역 시작점으로 복귀
                    if (ws.isPlaying()) {{
                        const curr = ws.getCurrentTime();
                        if (curr < region.start || curr > region.end) {{
                            ws.setTime(region.start);
                        }}
                    }}
                }});
                
                function updateUI(region) {{
                    const start = region.start.toFixed(2);
                    const end = region.end.toFixed(2);
                    
                    const regionTimeEl = document.getElementById('region-time');
                    if (regionTimeEl) {{
                        regionTimeEl.innerHTML = `<span class="time-span">${{start}}s ~ ${{end}}s</span>`;
                    }}
                    
                    // localStorage에 영역 정보 저장 (SET 버튼 릴레이용)
                    try {{
                        localStorage.setItem('wf_region_start', start);
                        localStorage.setItem('wf_region_end', end);
                    }} catch(e) {{ console.error('localStorage write failed:', e); }}
                    
                    const btnStart = document.getElementById('copyStartBtn');
                    const btnEnd = document.getElementById('copyEndBtn');
                    const btnBoth = document.getElementById('copyBothBtn');
                    
                    if (btnStart && btnEnd && btnBoth) {{
                        btnStart.disabled = false;
                        btnEnd.disabled = false;
                        btnBoth.disabled = false;
                        
                        btnStart.onclick = () => copyText(start, '시작 시간', false);
                        btnEnd.onclick = () => copyText(end, '종료 시간', false);
                        btnBoth.onclick = () => copyText(start + ', ' + end, '시작 및 종료 시간', true);
                    }}
                }}
                
                function copyText(text, type, storeForPaste) {{
                    // PASTE 릴레이용 localStorage 저장 (모두 복사인 경우만)
                    if (storeForPaste) {{
                        try {{ localStorage.setItem('wf_clipboard', text); }} catch(e) {{}}
                    }}
                    // 시스템 클립보드 복사
                    if (navigator.clipboard && navigator.clipboard.writeText) {{
                        navigator.clipboard.writeText(text).then(() => {{
                            showToast(`${{type}} 복사 완료: ${{text}}`);
                        }}).catch(() => {{ fallbackCopy(text, type); }});
                    }} else {{
                        fallbackCopy(text, type);
                    }}
                }}
                
                function fallbackCopy(text, type) {{
                    const textArea = document.createElement("textarea");
                    textArea.value = text;
                    textArea.style.position = "fixed";
                    document.body.appendChild(textArea);
                    textArea.focus();
                    textArea.select();
                    try {{
                        const successful = document.execCommand('copy');
                        if (successful) {{
                            showToast(`${{type}} 복사 완료: ${{text}}`);
                        }} else {{
                            alert('복사 실패');
                        }}
                    }} catch (err) {{
                        alert('복사 실패: ' + err);
                    }}
                    document.body.removeChild(textArea);
                }}
                
                function showToast(msg) {{
                    const toast = document.getElementById('toast');
                    toast.innerText = msg;
                    toast.classList.add('show');
                    setTimeout(() => {{
                        toast.classList.remove('show');
                    }}, 2000);
                }}
            </script>
            """
            st.components.v1.html(waveform_html, height=260)
            
            # 3. 편집 인터페이스 설계
            st.write("---")
            st.write("### ✂️ 문장별 시작/종료점 설정")

            if len(df_t9.columns) < 4:
                st.error("⚠️ 엑셀 파일에 적어도 4개의 열(A:교재명, B:단원명, C:파일명, D:영어문장)이 존재해야 합니다.")
            else:
                # ── 작업 대상 교재 및 단원 필터링 추가 ──────────────────────
                st.write("#### 🔍 작업 대상 필터링")
                col_filt1, col_filt2 = st.columns(2)
                
                book_col = df_t9.columns[0]
                books = sorted(list(df_t9[book_col].dropna().astype(str).unique()))
                with col_filt1:
                    selected_book = st.selectbox("📚 작업할 교재 선택", books, key="l_edit_sel_book")
                
                temp_df = df_t9[df_t9[book_col].astype(str) == selected_book]
                unit_col = df_t9.columns[1]
                units = sorted(list(temp_df[unit_col].dropna().astype(str).unique()))
                with col_filt2:
                    selected_unit = st.selectbox("📑 작업할 단원 선택", units, key="l_edit_sel_unit")
                
                filtered_df = df_t9[
                    (df_t9[book_col].astype(str) == selected_book) & 
                    (df_t9[unit_col].astype(str) == selected_unit)
                ]
                st.write("---")
                
                # 대량의 문장이 올라왔을 때 페이지가 너무 길어지는 문제를 방지하기 위해 10개씩 페이지네이션 지원
                items_per_page = 10
                total_items = len(filtered_df)
                total_pages = (total_items - 1) // items_per_page + 1 if total_items > 0 else 1

                col_page1, col_page2 = st.columns([1, 4])
                with col_page1:
                    page_num = st.number_input("페이지 번호", min_value=1, max_value=total_pages, value=1, step=1, key="l_edit_page_num")
                with col_page2:
                    st.write(f"선택한 교재/단원에 총 {total_items}개의 문장이 있습니다. (페이지: {page_num} / {total_pages})")

                start_idx = (page_num - 1) * items_per_page
                end_idx = min(start_idx + items_per_page, total_items)
                
                # 세션 상태에 시작/종료 값 초기화
                if "slices_state" not in st.session_state:
                    st.session_state.slices_state = {}

                st.write("")
                # 필터링된 데이터에서 해당 페이지 범위 슬라이싱
                page_df = filtered_df.iloc[start_idx:end_idx]
                
                for idx, row in page_df.iterrows():
                    val_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                    val_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else "Unassigned"
                    val_c = str(row.iloc[2]).strip() if not pd.isna(row.iloc[2]) else ""
                    val_d = str(row.iloc[3]).strip() if not pd.isna(row.iloc[3]) else ""

                    # 빈 줄은 스킵
                    if not val_d and not val_c:
                        continue

                    # 세션 상태값 가져오기/초기화
                    if idx not in st.session_state.slices_state:
                        st.session_state.slices_state[idx] = {"start": 0.0, "end": 0.0}

                    with st.container():
                        st.markdown(f"**[{idx + 1}] {val_d}** (저장 파일명: `{val_c}.mp3`) - {val_a} / {val_b}")
                        
                        col_num1, col_num2, col_control, col_listen = st.columns([2.5, 2.5, 2.0, 2.0])
                        with col_num1:
                            start_val = st.number_input(
                                f"시작(초) #{idx+1}", 
                                min_value=0.0, 
                                max_value=duration_sec, 
                                value=st.session_state.slices_state[idx]["start"], 
                                step=0.1, 
                                format="%.2f",
                                key=f"start_in_{idx}"
                            )
                            st.session_state.slices_state[idx]["start"] = start_val
                        with col_num2:
                            end_val = st.number_input(
                                f"종료(초) #{idx+1}", 
                                min_value=0.0, 
                                max_value=duration_sec, 
                                value=st.session_state.slices_state[idx]["end"] if st.session_state.slices_state[idx]["end"] > 0 else 0.0, 
                                step=0.1, 
                                format="%.2f",
                                key=f"end_in_{idx}"
                            )
                            st.session_state.slices_state[idx]["end"] = end_val
                        with col_control:
                            rendering_control_buttons(idx)
                        
                        with col_listen:
                            st.write("🔬 미리듣기")
                            # 시작 시각이 종료 시각보다 작은 유효한 입력일 때만 미리듣기 노출
                            if start_val < end_val:
                                if st.button(f"🔊 재생 #{idx+1}", key=f"btn_preview_{idx}", use_container_width=True):
                                    try:
                                        start_ms = int(start_val * 1000)
                                        end_ms = int(end_val * 1000)
                                        sliced_seg = audio_segment[start_ms:end_ms]
                                        
                                        preview_buf = io.BytesIO()
                                        sliced_seg.export(preview_buf, format="mp3")
                                        st.audio(preview_buf.getvalue(), format="audio/mp3")
                                    except Exception as ex_slice:
                                        st.error(f"미리듣기 생성 에러: {ex_slice}")
                            else:
                                st.caption("시작 < 종료여야 재생 가능")
                                
                    st.write("---")

                # 일괄 처리 영역
                st.write("### 📥 선택된 교재/단원 사운드 잘라내기 실행")
                st.write("페이지에 관계없이 **시작 및 종료 시간이 올바르게 입력된(시작 < 종료)** 모든 행의 오디오를 잘라내어 저장합니다.")
                
                if st.button("▶ 모든 사운드 끊어내기 및 저장 실행", type="primary", use_container_width=True, key="btn_slice_all_run"):
                    with st.spinner("사운드를 잘라내어 저장 폴더에 기록하는 중..."):
                        try:
                            saved_count = 0
                            progress_bar_all = st.progress(0)
                            status_text_all = st.empty()
                            total_rows_t9 = len(filtered_df)

                            for index, row in filtered_df.iterrows():
                                val_a = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                                val_b = str(row.iloc[1]).strip() if not pd.isna(row.iloc[1]) else "Unassigned"
                                val_c = str(row.iloc[2]).strip() if not pd.isna(row.iloc[2]) else ""
                                val_d = str(row.iloc[3]).strip() if not pd.isna(row.iloc[3]) else ""

                                if not val_c or not val_d:
                                    continue

                                # 세션 상태에서 시작/종료 값 획득
                                slice_info = st.session_state.slices_state.get(index, {"start": 0.0, "end": 0.0})
                                s_val = slice_info["start"]
                                e_val = slice_info["end"]

                                if s_val >= e_val or e_val <= 0:
                                    continue # 유효하지 않은 구간은 건너뜀

                                # 저장 경로 빌드
                                target_dir = os.path.join(
                                    base_output_dir_t9, 
                                    sanitize_filename(val_a) if val_a else "Unassigned", 
                                    sanitize_filename(val_b)
                                )
                                filename = sanitize_filename(val_c) + ".mp3"
                                file_path = os.path.join(target_dir, filename)

                                if not os.path.exists(target_dir):
                                    os.makedirs(target_dir, exist_ok=True)

                                # pydub 자르기 및 저장
                                s_ms = int(s_val * 1000)
                                e_ms = int(e_val * 1000)
                                cut_segment = audio_segment[s_ms:e_ms]
                                
                                # 저장 (덮어쓰기 지원)
                                cut_segment.export(file_path, format="mp3")
                                saved_count += 1

                                progress_bar_all.progress((index + 1) / total_rows_t9)
                                status_text_all.text(f"저장 중: {index + 1}/{total_rows_t9} (저장됨: {saved_count}개)")

                            st.success(f"🎉 사운드 끊어내기 및 저장 완료! 총 {saved_count}개의 파일이 생성 및 기록되었습니다.")
                            st.info(f"📂 저장 위치: `{os.path.abspath(base_output_dir_t9)}`")
                            
                            # ZIP 다운로드 지원 (현재 선택된 교재/단원 하위만 압축)
                            sub_dir_path = os.path.join(
                                sanitize_filename(selected_book) if selected_book else "Unassigned",
                                sanitize_filename(selected_unit)
                            )
                            zip_data = create_zip_of_directory(base_output_dir_t9, sub_directory=sub_dir_path)
                            if zip_data:
                                file_name_zip = f"{sanitize_filename(selected_unit)}_sliced.zip"
                                st.download_button(
                                    label=f"📥 선택된 단원({selected_unit}) 사운드 다운로드 (ZIP)",
                                    data=zip_data,
                                    file_name=file_name_zip,
                                    mime="application/zip",
                                    use_container_width=True,
                                    key="dl_zip_t9"
                                )
                        except Exception as e_run:
                            st.error(f"오디오 슬라이싱 실행 중 에러 발생: {e_run}")
        except Exception as e_init:
            st.error(f"파일 로드 및 전처리 오류: {e_init}")
